from flask import Flask, render_template, request, send_from_directory, url_for, redirect, send_file, flash, session
from markupsafe import escape
from datetime import datetime as dt
from num2words import num2words
import os, docx
from docx.shared import Cm
from werkzeug import exceptions


def folderMaker(path):
    try:
        os.mkdir(path)
    except:
        pass

_slash = "\\" if os.name == "nt" else "/"

owner = ["owner", "9990009990"]

SNAME = "000"
SADDRESS = "111"
PATH = os.getcwd()+f"{_slash}word"
folderMaker(PATH)
folderMaker(PATH+f"{_slash}bill{_slash}")
folderMaker(PATH+f"{_slash}receipt{_slash}")
LIST = []
def priceList():
    if "pricelist.docx" in os.listdir(os.getcwd()):
        LIST.clear()
        doc = docx.Document("pricelist.docx")
        i = 0
        if doc.tables:
            for row in doc.tables[0].rows:
                if row.cells[0].text.strip()!="":
                    check = [row.cells[0].text, row.cells[1].text, row.cells[2].text, i]
                    if check not in LIST:
                        LIST.append([row.cells[0].text, float(row.cells[1].text), float(row.cells[2].text), i]) 
                i+=1
    else:
        doc = docx.Document()
        doc.save("pricelist.docx")
    return LIST

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1000 * 1000
app.config["SECRET_KEY"] = "ZXZ$%^&*(0987654XYXZXZ"

@app.errorhandler(404)
def page_not_found(e):
    return render_template("404.html")

@app.errorhandler(exceptions.InternalServerError)
def internal_server_error(e):
    return redirect("/")

@app.before_request
def before_request():
    global SNAME, SADDRESS
    if "config.txt" in os.listdir(os.getcwd()):
        f = open("config.txt", "r").read()
        if "\n" in f:
            SNAME, SADDRESS = f.split("\n")[:2]
        else:
            pass
    if "__uname__" not in session and "__psswd__" not in session and  "login" not in request.url:
            return redirect("/login")
    if "__uname__" in session and "__psswd__" in session:
        if (session["__uname__"] in owner and session["__psswd__"] in owner) or (f'{session["__uname__"]},' + f' {session["__psswd__"]}' in open("users.log", 'r').read()):
            if session["__uname__"] not in owner and session["__psswd__"] not in owner:
                if request.url == "http://127.0.0.1:5000/files" or "logout" in request.url or "static" in request.url or request.url.endswith("download"):
                    pass
                else:
                    return redirect("/files")
        else:
            session.pop("__uname__", None)
            session.pop("__psswd__", None)
            flash("invalid username or password")
            return redirect("/login")

@app.route("/login", methods=["POST", "GET"])
def login():
    if request.method == "POST":
        if request.form["email"] in owner and request.form['password'] in owner:
            session["__uname__"] = request.form["email"]
            session["__psswd__"] = request.form["password"]
            return redirect("/")
        if "users.log" in os.listdir("."):
            for i in open("users.log", "r").read().split("\n"):
                if ", " in i:
                    __user, __psswd = i.split(", ")
                    if __user == request.form["email"] and __psswd == request.form['password']:
                        session["__uname__"] = request.form["email"]
                        session["__psswd__"] = request.form["password"]
                        return redirect("/")
            else:
                flash("Invalid username or password")
                return redirect("/login")
        else:
            flash("Invalid username or password")
            return redirect("/login")
    
    return render_template("login.html")

@app.route("/logout")
def logout():
    if "__uname__" not in session and "__psswd__" not in session:
        return redirect("/login")
    session.pop("__uname__", None)
    session.pop("__psswd__", None)
    return redirect("/login")

@app.route("/adduser", methods = ["POST", "GET"])
def signup():
    if request.method == "POST":
        f = open("users.log", "a")
        f.write(f"{request.form['email']}, {request.form['password']}\n")
        f.close()
        return redirect("/")
    return render_template("signup.html")

@app.route("/config", methods = ["POST", "GET"])
def configure():
    global SNAME, SADDRESS
    if request.method == "POST":
        SNAME = request.form["SNAME"].strip()
        SADDRESS = request.form["SADDRESS"].strip()
        f = open("config.txt", "w")
        f.write(f"{SNAME}\n{SADDRESS}")
        f.close()
        return redirect("/")
    return render_template("config.html")

@app.route("/")
@app.route("/home")
def i():
    return render_template("index.html", SNAME=SNAME, SADDRESS=SADDRESS)

@app.route("/inventory", methods=["POST", "GET"])
def pricelist():
    if request.method == "POST":
        try:
            doc = docx.Document("pricelist.docx")
        except:
            doc = docx.Document()
        if "index" in request.form:
            table = doc.tables[0]
            row = table.rows[int(request.form["index"])]
            row.cells[0].text = request.form["product"]
            row.cells[1].text = request.form["quantity"]
            row.cells[2].text = request.form["price"]
            doc.save("pricelist.docx")
            flash("Edit Success.")
            return redirect("/inventory")
        try:
            if doc.tables:
                table = doc.tables[0]
                new_row = table.add_row()
                new_row.cells[0].text = request.form["product"]
                new_row.cells[1].text = request.form['quantity']
                new_row.cells[2].text = request.form["price"]
            else:
                table = doc.add_table(rows = 1, cols = 3)
                table.style = "Table Grid"
                table.rows[0].cells[0].text = request.form["product"]
                table.rows[0].cells[1].text = request.form['quantity']
                table.rows[0].cells[2].text = request.form["price"]
            doc.save("pricelist.docx")
        except PermissionError:
            flash("Failed. Please Try After Closing The File.")
            return redirect("/inventory")
        flash("Added Successlully")
        return redirect("/inventory")
    return render_template("priceList.html", SNAME=SNAME, SADDRESS=SADDRESS, LIST = priceList())

@app.route("/deleterow", methods=["POST", "GET"])
def deleterow():
    if request.method == "POST":
        doc = docx.Document("pricelist.docx")
        doc.tables[0]._tbl.remove(doc.tables[0].rows[int(request.form["index"])]._tr)
        try:
            doc.save("pricelist.docx")
        except PermissionError:
            flash("Failed. Please Try After Closing The File.")
            return redirect("/inventory")
        flash("Removed Successfully")
        return redirect("/inventory")

@app.route("/bill")
def h():
    bNo = len(list(set([i.split(" ")[0] for i in os.listdir(PATH+f"{_slash}bill{_slash}") if i.lower().endswith(".docx") and i.upper().startswith("B")])))
    start = 1
    tf = False
    while tf == False:
        for i in os.listdir(PATH+f"{_slash}bill{_slash}"):
            if i.startswith("B"+str(start)):
                start+=1
                break
        else:
            break
    return render_template("bill.html", SNAME=SNAME, SADDRESS=SADDRESS, LIST=priceList(), bNo = start)

@app.route("/bill/edit/<string>")
def edit_bill(string):
    bNo, _name = string.split(" ")[0:2]
    bNo = bNo[1:]
    doc = docx.Document(PATH+f"{_slash}bill{_slash}"+string)
    l = []
    if doc.tables:
        r = doc.tables[0].rows[1:-1]
        total_amount = doc.tables[0].rows[-1].cells[-1].text
        for i in r:
            l.append([i.cells[0].text, i.cells[1].text, i.cells[2].text, i.cells[3].text, i.cells[4].text])
    return render_template("bill.html", SNAME=SNAME, SADDRESS=SADDRESS, LIST=priceList(), bNo = bNo, l = l, total_amount = total_amount, _name = _name)

@app.route("/receipt")
def r():
    receiptNo = len(list(set([i.split(" ")[0] for i in os.listdir(PATH+f"{_slash}bill{_slash}") if i.lower().endswith(".docx") and i.upper().startswith("R")])))
    start = 1
    tf = False
    while tf == False:
        for i in os.listdir(PATH+f"{_slash}receipt{_slash}"):
            if i.startswith("R"+str(start)):
                start+=1
                break
        else:
            break
    return render_template("receipt.html", SNAME=SNAME, SADDRESS=SADDRESS, receiptNo = start)

@app.route("/submitReceipt", methods=["POST", "GET"])
def sr():
    if request.method == "POST":
        name, address, amount, receiptNo = request.form["name"], request.form["address"], request.form["amount"], request.form["receiptNo"]
        doc = docx.Document()
        head = doc.add_heading(f"{SNAME}\n{SADDRESS}", 0)
        head.alignment = 1
        doc.add_heading(f"Receipt No: - {receiptNo}")
        doc.add_paragraph(f"Date: - {dt.now().strftime('%d / %m / %Y')}")
        doc.add_paragraph(f"Name: - {name}")
        if address:
            doc.add_paragraph(f"Address: - {address}")
        doc.add_paragraph(f"Amount (in words): - {num2words(amount)}")
        doc.add_paragraph(f"Amount: - {amount}")
        for i in os.listdir(PATH+f"{_slash}receipt{_slash}"):
            if receiptNo+" " in i:
                os.remove(PATH+f"{_slash}receipt{_slash}"+i)
        file_name = f"R{receiptNo} {name} {address} {dt.now().strftime('%d_%m_%Y')}".strip()
        doc.save(PATH+f"{_slash}receipt{_slash}"+file_name+".docx")
        return redirect("/files/"+file_name+".docx")

@app.route("/print", methods=["POST", "GET"])
def generateDocx():
    if request.method == "POST":
        data = request.form["data"]
        na, data = data.split("\r\n\r\n")
        name, address, billNo = na.split("==_==")
        bPath = PATH+f"{_slash}bill"
        c = []
        s = 0
        doc = docx.Document()
        head = doc.add_heading(f"{SNAME}\n{SADDRESS}", 0)
        head.alignment = 1
        doc.add_paragraph(f"{billNo}")
        doc.add_paragraph(f"Date: - {dt.now().strftime('%d / %m / %Y')}" if "edit" not in request.referrer else f"Date: - {request.referrer.split('/')[-1].replace('%20', ' ').split('.')[0]}")
        if name.strip() and name!="NAME":
            doc.add_paragraph(f"Name: - {name}")
        else:
            name = "Cash"
            doc.add_paragraph(f"Name: - Cash")
        if address.strip() and address!="ADDRESS":
            doc.add_paragraph(f"Address: - {address}")
        else:
            address = ""
        t = doc.add_table(rows = 1, cols = 5)
        t.style = "Table Grid"
        t.rows[0].cells[0].text = "Sl No."
        t.rows[0].cells[0].width = Cm(2)
        t.rows[0].cells[1].text = "Particular"
        t.rows[0].cells[1].width = Cm(7)
        t.rows[0].cells[2].text = "Quantity"
        t.rows[0].cells[2].width = Cm(2)
        t.rows[0].cells[3].text = "Price"
        t.rows[0].cells[3].width = Cm(2)
        t.rows[0].cells[4].text = "Total"
        t.rows[0].cells[4].width = Cm(2)
        for i in data.split("\r\n"):
            if i!="":
                new_row = t.add_row()
                slNo, particular, quantity, price, total = i.split("==_==")
                new_row.cells[0].text = slNo
                new_row.cells[0].width = Cm(2)
                new_row.cells[1].text = particular
                new_row.cells[1].width = Cm(7)
                _pdoc = docx.Document("pricelist.docx")
                if _pdoc.tables:
                    _ptable = _pdoc.tables[0]
                    for _rows in _ptable.rows:
                        if _rows.cells[0].text.lower() == particular.lower():
                            _rows.cells[1].text = str(float(_rows.cells[1].text) - float(quantity))
                            # if float(_rows.cells[1].text)<=0:
                            #     quantity = str(float(quantity) + float(_rows.cells[1].text))
                            #     total = str(float(quantity)*float(price))
                            #     _pdoc.tables[0]._tbl.remove(_rows._tr)
                    _pdoc.save("pricelist.docx")
                new_row.cells[2].text = quantity
                new_row.cells[2].width = Cm(2)
                new_row.cells[3].text = price
                new_row.cells[3].width = Cm(2)
                new_row.cells[4].text = total
                new_row.cells[4].width = Cm(2)
                c.append(i.split("==_=="))
                s+=float(total)
        new_row = t.add_row()
        new_row.cells[3].text = "Total"
        new_row.cells[4].text = str(s)
        billNo = billNo.split(": - ")[1]
        if "edit" not in request.referrer:
            fileName = f"{bPath}{_slash}B{billNo} {name} {address} {dt.now().strftime('%d_%m_%Y')}".strip() +".docx"
        else:
            fileName = f"{bPath}{_slash}{request.referrer.split('/')[-1].replace('%20', ' ')}"
            print(fileName)
        for i in os.listdir(bPath):
            if billNo+" " in i:
                os.remove(bPath+f"{_slash}"+i)
        doc.save(fileName)
        return redirect(f"/files/B{billNo} {name} {address} {dt.now().strftime('%d_%m_%Y')}".strip()+".docx")

@app.route("/delete", methods=["POST", "GET"])
def delete_file():
    if request.method == "POST":
        n = request.form["_name"]
        try:    
            if n.startswith("B"):
                os.remove(os.getcwd()+f"{_slash}word{_slash}bill{_slash}"+n)
            elif n.startswith("R"):
                os.remove(os.getcwd()+f"{_slash}word{_slash}receipt{_slash}"+n)
            flash("Deleted "+n+" Successfully.")
        except:
            flash("Unable to delete the file.")
        return redirect("/")

@app.route("/files", methods=["POST", "GET"])
def see():
    if session["__uname__"] == owner[0] and session["__psswd__"] == owner[1]:
        if request.method == "POST":
            term = (request.form["search"]).lower()
            files = []
            for i in os.listdir(PATH+f"{_slash}bill{_slash}"):
                if i.endswith(".docx") and term in i.lower():
                    doc = docx.Document(PATH+f"{_slash}bill{_slash}"+i)
                    files.append([i, round((os.path.getsize(PATH+f"{_slash}bill{_slash}"+i)//1024)/1024, 3), doc.tables[0].rows[-1].cells[-1].text])
            for i in os.listdir(PATH+f"{_slash}receipt{_slash}"):
                if i.endswith(".docx") and term in i.lower():
                    doc = docx.Document(PATH+f"{_slash}receipt{_slash}"+i)
                    files.append([i, (os.path.getsize(PATH+f"{_slash}receipt{_slash}"+i)//1024)/1024, doc.paragraphs[-1].text.split(": - ")[-1]])
            return render_template("files.html", SNAME=SNAME, SADDRESS=SADDRESS, files=files)
        files = []
        for i in os.listdir(PATH+f"{_slash}bill{_slash}"):
            if i.endswith(".docx"):
                doc = docx.Document(PATH+f"{_slash}bill{_slash}"+i)
                files.append([i, (os.path.getsize(PATH+f"{_slash}bill{_slash}"+i)//1024)/1024, doc.tables[0].rows[-1].cells[-1].text])
        for i in os.listdir(PATH+f"{_slash}receipt{_slash}"):
            if i.endswith(".docx"):
                doc = docx.Document(PATH+f"{_slash}receipt{_slash}"+i)
                files.append([i, (os.path.getsize(PATH+f"{_slash}receipt{_slash}"+i)//1024)/1024, doc.paragraphs[-1].text.split(": - ")[-1]])
        return render_template("files.html", SNAME=SNAME, SADDRESS=SADDRESS, files=files)
    else:
        term = session["__uname__"].lower()
        files = []
        for i in os.listdir(PATH+f"{_slash}bill{_slash}"):
            if i.endswith(".docx") and term in i.lower():
                doc = docx.Document(PATH+f"{_slash}bill{_slash}"+i)
                files.append([i, round((os.path.getsize(PATH+f"{_slash}bill{_slash}"+i)//1024)/1024, 3), doc.tables[0].rows[-1].cells[-1].text])
        for i in os.listdir(PATH+f"{_slash}receipt{_slash}"):
            if i.endswith(".docx") and term in i.lower():
                doc = docx.Document(PATH+f"{_slash}receipt{_slash}"+i)
                files.append([i, (os.path.getsize(PATH+f"{_slash}receipt{_slash}"+i)//1024)/1024, doc.paragraphs[-1].text.split(": - ")[-1]])
        return render_template("files.html", SNAME=SNAME, SADDRESS=SADDRESS, files=files)

@app.route("/files/<path:path>")
def watch(path):
    if path.startswith("B"):
        doc = docx.Document(PATH+f"{_slash}bill{_slash}"+path)
        return render_template("watch.html", SNAME=SNAME, SADDRESS=SADDRESS, rows = doc.tables[0].rows[1:], para = doc.paragraphs[1:], name=path)
    elif path.startswith("R"):
        doc = docx.Document(PATH+"/receipt/"+path)
        return render_template("watch.html", SNAME=SNAME, SADDRESS=SADDRESS, para = doc.paragraphs[1:], name=path)

@app.route("/files/<path:path>/download")
def download(path):
    if path.startswith("B"):
        return send_file(PATH+f"{_slash}bill{_slash}"+path, as_attachment=True)
    else:
        return send_file(PATH+f"{_slash}receipt{_slash}"+path, as_attachment=True)