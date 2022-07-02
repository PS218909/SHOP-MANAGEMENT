from flask import Flask, render_template, request, send_from_directory, url_for, redirect, send_file, flash
from markupsafe import escape
from datetime import datetime as dt
from num2words import num2words
import os, docx, webbrowser
from docx.shared import Cm

def folderMaker(path):
    try:
        os.mkdir(path)
    except:
        pass

PATH = os.getcwd()+"\\word"
folderMaker(PATH)
folderMaker(PATH+"\\bill\\")
folderMaker(PATH+"\\receipt\\")
LIST = []
def priceList():
    if "pricelist.docx" in os.listdir(os.getcwd()):
        LIST.clear()
        doc = docx.Document("pricelist.docx")
        i = 0
        if doc.tables:
            for row in doc.tables[0].rows:
                if row.cells[0].text.strip()!="":
                    check = [row.cells[0].text, row.cells[1].text, row.cells[2].text, i]
                    if check not in LIST:
                        LIST.append([row.cells[0].text, int(row.cells[1].text), int(row.cells[2].text), i]) 
                i+=1
    else:
        doc = docx.Document()
        doc.save("pricelist.docx")
    return LIST

SNAME = "XYZ"
SADDRESS = "ZYX"

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1000 * 1000
app.config["SECRET_KEY"] = "ZXZXYXZXZ"

@app.errorhandler(404)
def page_not_found(e):
    return render_template("404.html")

@app.route("/")
@app.route("/home")
def i():
    return render_template("index.html", SNAME=SNAME, SADDRESS=SADDRESS)

@app.route("/pricelist", methods=["POST", "GET"])
def pricelist():
    if request.method == "POST":
        try:
            doc = docx.Document("pricelist.docx")
        except:
            doc = docx.Document()
        try:
            if doc.tables:
                table = doc.tables[0]
                new_row = table.add_row()
                new_row.cells[0].text = request.form["product"]
                new_row.cells[1].text = request.form['quantity']
                new_row.cells[2].text = request.form["price"]
            else:
                table = doc.add_table(rows = 1, cols = 3)
                table.style = "Table Grid"
                table.rows[0].cells[0].text = request.form["product"]
                table.rows[0].cells[1].text = request.form['quantity']
                table.rows[0].cells[2].text = request.form["price"]
            doc.save("pricelist.docx")
        except PermissionError:
            flash("Failed. Please Try After Closing The File.")
            return redirect("/pricelist")
        flash("Added Successlully")
        return redirect("/pricelist")
    return render_template("priceList.html", SNAME=SNAME, SADDRESS=SADDRESS, LIST = priceList())

@app.route("/deleterow", methods=["POST", "GET"])
def deleterow():
    if request.method == "POST":
        doc = docx.Document("pricelist.docx")
        doc.tables[0]._tbl.remove(doc.tables[0].rows[int(request.form["index"])]._tr)
        try:
            doc.save("pricelist.docx")
        except PermissionError:
            flash("Failed. Please Try After Closing The File.")
            return redirect("/pricelist")
        flash("Removed Successfully")
        return redirect("/pricelist")

@app.route("/bill")
def h():
    bNo = len(list(set([i.split(" ")[0] for i in os.listdir(PATH+"\\bill\\") if i.lower().endswith(".docx") and i.upper().startswith("B")])))
    start = 1
    tf = False
    while tf == False:
        for i in os.listdir(PATH+"\\bill\\"):
            if i.startswith("B"+str(start)):
                start+=1
                break
        else:
            break
    return render_template("bill.html", SNAME=SNAME, SADDRESS=SADDRESS, LIST=priceList(), bNo = start)

@app.route("/receipt")
def r():
    receiptNo = len(list(set([i.split(" ")[0] for i in os.listdir(PATH+"\\bill\\") if i.lower().endswith(".docx") and i.upper().startswith("R")])))
    start = 1
    tf = False
    while tf == False:
        for i in os.listdir(PATH+"\\receipt\\"):
            if i.startswith("R"+str(start)):
                start+=1
                break
        else:
            break
    return render_template("receipt.html", SNAME=SNAME, SADDRESS=SADDRESS, receiptNo = start)

@app.route("/submitReceipt", methods=["POST", "GET"])
def sr():
    if request.method == "POST":
        name, address, amount, receiptNo = request.form["name"], request.form["address"], request.form["amount"], request.form["receiptNo"]
        doc = docx.Document()
        head = doc.add_heading(f"{SNAME}\n{SADDRESS}", 0)
        head.alignment = 1
        doc.add_heading(f"Receipt No: - {receiptNo}")
        doc.add_paragraph(f"Date: - {dt.now().strftime('%d / %m / %Y')}")
        doc.add_paragraph(f"Name: - {name}")
        if address:
            doc.add_paragraph(f"Address: - {address}")
        doc.add_paragraph(f"Amount (in words): - {num2words(amount)}")
        doc.add_paragraph(f"Amount: - {amount}")
        for i in os.listdir(PATH+"\\receipt\\"):
            if receiptNo+" " in i:
                os.remove(PATH+"\\receipt\\"+i)
        doc.save(PATH+"\\receipt\\"+f"R{receiptNo} {name} {address}".strip()+".docx")
        f = open("data.csv", 'a')
        f.write(f"R{receiptNo}, {name}, {address}, {amount}\n")
        f.close()
        return redirect("/files/"+f"R{receiptNo} {name} {address}".strip()+".docx")

@app.route("/print", methods=["POST", "GET"])
def generateDocx():
    if request.method == "POST":
        data = request.form["data"]
        na, data = data.split("\r\n\r\n")
        name, address, billNo = na.split("==_==")
        bPath = PATH+"\\bill"
        c = []
        s = 0
        doc = docx.Document()
        head = doc.add_heading(f"{SNAME}\n{SADDRESS}", 0)
        head.alignment = 1
        doc.add_paragraph(f"{billNo}")
        doc.add_paragraph(f"Date: - {dt.now().strftime('%d / %m / %Y')}")
        if name.strip() and name!="NAME":
            doc.add_paragraph(f"Name: - {name}")
        else:
            name = "Cash"
            doc.add_paragraph(f"Name: - Cash")
        if address.strip() and address!="ADDRESS":
            doc.add_paragraph(f"Address: - {address}")
        else:
            address = ""
        t = doc.add_table(rows = 1, cols = 5)
        t.style = "Table Grid"
        t.rows[0].cells[0].text = "Sl No."
        t.rows[0].cells[0].width = Cm(2)
        t.rows[0].cells[1].text = "Particular"
        t.rows[0].cells[1].width = Cm(7)
        t.rows[0].cells[2].text = "Quantity"
        t.rows[0].cells[2].width = Cm(2)
        t.rows[0].cells[3].text = "Price"
        t.rows[0].cells[3].width = Cm(2)
        t.rows[0].cells[4].text = "Total"
        t.rows[0].cells[4].width = Cm(2)
        for i in data.split("\r\n"):
            if i!="":
                new_row = t.add_row()
                slNo, particular, quantity, price, total = i.split("==_==")
                new_row.cells[0].text = slNo
                new_row.cells[0].width = Cm(2)
                new_row.cells[1].text = particular
                new_row.cells[1].width = Cm(7)
                new_row.cells[2].text = quantity
                new_row.cells[2].width = Cm(2)
                _pdoc = docx.Document("pricelist.docx")
                if _pdoc.tables:
                    _ptable = _pdoc.tables[0]
                    for _rows in _ptable.rows:
                        if _rows.cells[0].text.lower() == particular.lower():
                            _rows.cells[1].text = str(int(_rows.cells[1].text) - int(quantity))
                            if int(_rows.cells[1].text)<=0:
                                _pdoc.tables[0]._tbl.remove(_rows._tr)
                    _pdoc.save("pricelist.docx")
                new_row.cells[3].text = price
                new_row.cells[3].width = Cm(2)
                new_row.cells[4].text = total
                new_row.cells[4].width = Cm(2)
                c.append(i.split("==_=="))
                s+=int(total)
        new_row = t.add_row()
        new_row.cells[3].text = "Total"
        new_row.cells[4].text = str(s)
        billNo = billNo.split(": - ")[1]
        fileName = f"{bPath}\\B{billNo} {name} {address}".strip() +".docx"
        for i in os.listdir(bPath):
            if billNo+" " in i:
                os.remove(bPath+"\\"+i)
        doc.save(fileName)
        f = open("data.csv", "a")
        f.write(f"{str(dt.datetime.now()).split(' ')[0]}, B{billNo}, {name}, {address}, {total}\n")
        f.close()
        return redirect(f"/files/B{billNo} {name} {address}".strip()+".docx")

@app.route("/files", methods=["POST", "GET"])
def see():
    if request.method == "POST":
        term = (request.form["search"]).lower()
        files = []
        for i in os.listdir(PATH+"\\bill\\"):
            if i.endswith(".docx") and term in i.lower():
                doc = docx.Document(PATH+"\\bill\\"+i)
                files.append([i, (os.path.getsize(PATH+"\\bill\\"+i)//1024)/1024, doc.tables[0].rows[-1].cells[-1].text])
        for i in os.listdir(PATH+"\\receipt\\"):
            if i.endswith(".docx") and term in i.lower():
                doc = docx.Document(PATH+"\\receipt\\"+i)
                files.append([i, (os.path.getsize(PATH+"\\receipt\\"+i)//1024)/1024, doc.paragraphs[-1].text.split(": - ")[-1]])
        return render_template("files.html", SNAME=SNAME, SADDRESS=SADDRESS, files=files)
    files = []
    for i in os.listdir(PATH+"\\bill\\"):
        if i.endswith(".docx"):
            doc = docx.Document(PATH+"\\bill\\"+i)
            files.append([i, (os.path.getsize(PATH+"\\bill\\"+i)//1024)/1024, doc.tables[0].rows[-1].cells[-1].text])
    for i in os.listdir(PATH+"\\receipt\\"):
        if i.endswith(".docx"):
            doc = docx.Document(PATH+"\\receipt\\"+i)
            files.append([i, (os.path.getsize(PATH+"\\receipt\\"+i)//1024)/1024, doc.paragraphs[-1].text.split(": - ")[-1]])
    return render_template("files.html", SNAME=SNAME, SADDRESS=SADDRESS, files=files)

@app.route("/files/<path:path>")
def watch(path):
    if path.startswith("B"):
        doc = docx.Document(PATH+"\\bill\\"+path)
        return render_template("watch.html", SNAME=SNAME, SADDRESS=SADDRESS, rows = doc.tables[0].rows[1:], para = doc.paragraphs[1:], name=path)
    elif path.startswith("R"):
        doc = docx.Document(PATH+"/receipt/"+path)
        return render_template("watch.html", SNAME=SNAME, SADDRESS=SADDRESS, para = doc.paragraphs[1:], name=path)

@app.route("/files/<path:path>/download")
def download(path):
    return send_file(PATH+"\\bill\\"+path, as_attachment=True)

webbrowser.open('http://127.0.0.1:5000')
app.run()